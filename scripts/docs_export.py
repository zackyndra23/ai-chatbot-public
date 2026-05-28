"""
docs_export.py — generate Word (.docx) exports from the canonical markdown in docs/.

Usage:
    python scripts/docs_export.py --section all
    python scripts/docs_export.py --section architecture
    python scripts/docs_export.py --section modules
    python scripts/docs_export.py --section api
    python scripts/docs_export.py --section ops

Outputs to docs/exports/ (gitignored).

Uses pandoc if available (best fidelity), falls back to python-docx via
markdown → HTML → docx conversion.

This script is tooling, not application code — no dependency on core/app_config.
"""
from __future__ import annotations

import argparse
import shutil
import subprocess
import sys
from datetime import datetime
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = REPO_ROOT / "docs"
EXPORT_DIR = DOCS_DIR / "exports"

# Section definitions: order matters for the "all" export.
SECTIONS: dict[str, list[Path]] = {
    "architecture": [
        DOCS_DIR / "README.md",
        DOCS_DIR / "ARCHITECTURE.md",
    ],
    "modules": sorted((DOCS_DIR / "modules").rglob("*.md")),
    "api": sorted((DOCS_DIR / "api").glob("*.md")),
    "ops": sorted((DOCS_DIR / "ops").glob("*.md")),
}


def _git_sha() -> str:
    try:
        out = subprocess.run(
            ["git", "rev-parse", "--short", "HEAD"],
            cwd=REPO_ROOT,
            capture_output=True,
            text=True,
            check=True,
        )
        return out.stdout.strip()
    except Exception:
        return "unknown"


def _cover_page(title: str) -> str:
    return (
        f"% {title}\n"
        f"% RAG Chatbot v01 — Documentation Export\n"
        f"% Generated {datetime.now().strftime('%Y-%m-%d %H:%M')} "
        f"(commit: {_git_sha()})\n\n"
    )


def _gather(section: str) -> list[Path]:
    if section == "all":
        # Preserve logical order: overview → architecture → modules → api → ops.
        paths: list[Path] = []
        for key in ("architecture", "modules", "api", "ops"):
            paths.extend(SECTIONS[key])
        return paths
    if section not in SECTIONS:
        raise SystemExit(
            f"Unknown section: {section!r}. "
            f"Choose one of: all, {', '.join(SECTIONS)}"
        )
    return list(SECTIONS[section])


def _concat_markdown(paths: list[Path], title: str) -> str:
    parts = [_cover_page(title)]
    for p in paths:
        if not p.exists():
            continue
        rel = p.relative_to(REPO_ROOT).as_posix()
        parts.append(f"\n\n---\n\n# `{rel}`\n\n")
        parts.append(p.read_text(encoding="utf-8"))
    return "".join(parts)


def _have_pandoc() -> bool:
    return shutil.which("pandoc") is not None


def _export_pandoc(md_text: str, out_path: Path) -> None:
    """Preferred path: pandoc handles TOC, styling, and structure best."""
    tmp_md = out_path.with_suffix(".tmp.md")
    tmp_md.write_text(md_text, encoding="utf-8")
    try:
        subprocess.run(
            [
                "pandoc",
                str(tmp_md),
                "-o",
                str(out_path),
                "--toc",
                "--toc-depth=3",
                "--standalone",
                "-f",
                "gfm",
                "-t",
                "docx",
            ],
            check=True,
        )
    finally:
        try:
            tmp_md.unlink()
        except Exception:
            pass


def _export_python_docx(md_text: str, out_path: Path) -> None:
    """Fallback path: pure-Python, less pretty, no TOC."""
    try:
        import markdown  # type: ignore
        from docx import Document  # type: ignore
        from docx.shared import Pt  # type: ignore
        from bs4 import BeautifulSoup  # type: ignore
    except ImportError as e:
        raise SystemExit(
            f"Fallback export requires: markdown, python-docx, beautifulsoup4. "
            f"Install with: pip install markdown python-docx beautifulsoup4\n"
            f"Original error: {e}"
        )

    html = markdown.markdown(
        md_text,
        extensions=["fenced_code", "tables", "toc"],
    )
    soup = BeautifulSoup(html, "html.parser")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for element in soup.descendants:
        if element.name is None or not element.name:
            continue
        text = element.get_text(strip=False)
        if element.name in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            level = int(element.name[1])
            doc.add_heading(text.strip(), level=min(level, 4))
        elif element.name == "p":
            doc.add_paragraph(text.strip())
        elif element.name == "li":
            doc.add_paragraph(text.strip(), style="List Bullet")
        elif element.name == "pre":
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = "Consolas"
            run.font.size = Pt(9)

    doc.save(str(out_path))


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--section",
        default="all",
        choices=["all", *SECTIONS.keys()],
        help="Which section(s) to export.",
    )
    parser.add_argument(
        "--out",
        default=None,
        help="Override output path (default: docs/exports/<name>_<YYYYMMDD>.docx).",
    )
    parser.add_argument(
        "--force-fallback",
        action="store_true",
        help="Skip pandoc even if installed; use python-docx fallback.",
    )
    args = parser.parse_args()

    EXPORT_DIR.mkdir(parents=True, exist_ok=True)

    paths = _gather(args.section)
    if not paths:
        raise SystemExit(f"No markdown files found for section {args.section!r}.")

    title_map = {
        "all": "Full Documentation",
        "architecture": "Architecture Overview",
        "modules": "Module Reference",
        "api": "API Reference",
        "ops": "Operations Guide",
    }
    title = title_map.get(args.section, args.section.title())

    date_tag = datetime.now().strftime("%Y%m%d")
    default_name = f"rag_chatbot_v01_{args.section}_{date_tag}.docx"
    out_path = Path(args.out) if args.out else EXPORT_DIR / default_name

    md_text = _concat_markdown(paths, title)

    if _have_pandoc() and not args.force_fallback:
        print(f"[docs_export] Using pandoc → {out_path}")
        _export_pandoc(md_text, out_path)
    else:
        reason = "--force-fallback" if args.force_fallback else "pandoc not installed"
        print(f"[docs_export] {reason}; falling back to python-docx → {out_path}")
        _export_python_docx(md_text, out_path)

    print(
        f"[docs_export] done: {out_path.relative_to(REPO_ROOT).as_posix()} "
        f"({out_path.stat().st_size // 1024} KB)"
    )


if __name__ == "__main__":
    try:
        main()
    except SystemExit:
        raise
    except Exception as e:  # pragma: no cover
        print(f"[docs_export] error: {e}", file=sys.stderr)
        sys.exit(1)
